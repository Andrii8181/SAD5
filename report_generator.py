import docx
from docx.shared import Inches
import matplotlib.pyplot as plt
from tkinter import filedialog, messagebox

class ReportGenerator:
    """
    Клас для генерації звіту в форматі .docx.
    """
    def __init__(self, data_frame, analysis_type, analysis_results):
        self.df = data_frame
        self.analysis_type = analysis_type
        self.analysis_results = analysis_results

    def generate(self):
        """Створює та зберігає звіт."""
        try:
            doc = docx.Document()
            doc.add_heading(f"Звіт по аналізу даних: {self.analysis_type}", level=1)
            
            doc.add_heading("Початкові дані", level=2)
            doc.add_paragraph(self.df.to_string())

            doc.add_heading("Результати аналізу", level=2)
            if isinstance(self.analysis_results, str):
                doc.add_paragraph(self.analysis_results)
            else:
                doc.add_paragraph(self.analysis_results.to_string())

            doc.add_heading("Графіки", level=2)
            self.create_box_plot(doc)
            self.create_histogram(doc)

            doc_path = filedialog.asksaveasfilename(defaultextension=".docx", filetypes=[("Word Documents", "*.docx")])
            if doc_path:
                doc.save(doc_path)
                messagebox.showinfo("Успіх", f"Звіт збережено у: {doc_path}")

        except Exception as e:
            messagebox.showerror("Помилка", f"Не вдалося згенерувати звіт: {e}")

    def create_box_plot(self, doc):
        """Створює та додає боксплот до звіту."""
        try:
            plt.figure()
            self.df.boxplot()
            plt.title("Боксплот")
            plt.ylabel("Значення")
            plt.xlabel("Показники")
            plt.tight_layout()
            plt.savefig("boxplot.png")
            doc.add_paragraph("Боксплот:")
            doc.add_picture("boxplot.png", width=Inches(6))
        except Exception as e:
            doc.add_paragraph(f"Помилка при створенні боксплоту: {e}")

    def create_histogram(self, doc):
        """Створює та додає гістограму до звіту."""
        try:
            plt.figure()
            self.df.iloc[:, 0].hist()
            plt.title("Гістограма")
            plt.xlabel("Значення")
            plt.ylabel("Частота")
            plt.tight_layout()
            plt.savefig("histogram.png")
            doc.add_paragraph("Гістограма:")
            doc.add_picture("histogram.png", width=Inches(6))
        except Exception as e:
            doc.add_paragraph(f"Помилка при створенні гістограми: {e}")
